from random import randint
from random import shuffle
from docx import Document
from docx.shared import Pt,RGBColor,Cm
import docx
import os
from cn2an import an2cn

def answer(days,subject):
    try:
        d = Document('answer.docx')
    except:
        d = Document()
    a = d.add_paragraph()
    day = a.add_run('第'+str(days)+'天：')
    day.font.size = Pt(9)
    day.bold = True
    t = ''
    for x in range(len(subject)):
        e = str(eval(str(subject[x]).replace('=','')))
        t = t+'（'+str(x+1)+')'+e
    text = d.add_paragraph()
    text = text.add_run(t)
    text.font.size = Pt(9)
    d.save('answer.docx')



def blend():
    blend = []
    #a+b+c
    for x in range(10):
        while True: #Duplicate removal
            a = randint(15, 100)
            b = randint(15, 100)
            c = randint(15, 100)
            d = str(a)+'+'+str(b)+'+'+str(c)+'='
            if d in blend:
                pass
            else:
                blend.append(d)
                break

    #a+b-c
    for x in range(10):
        while True: #Duplicate removal
            a = randint(15, 100)
            b = randint(15, 100)
            c = randint(15, 100)
            d = str(a)+'+'+str(b)+'-'+str(c)+'='
            if d in blend or a+b<c:
                pass
            else:
                blend.append(d)
                break

    #a-b+c
    for x in range(10):
        while True: #Duplicate removal
            a = randint(15, 100)
            b = randint(15, 100)
            c = randint(15, 100)
            d = str(a)+'-'+str(b)+'+'+str(c)+'='
            if d in blend or a+c<b:
                pass
            else:
                blend.append(d)
                break
    #a-b-c
    for x in range(10):
        while True:  # Duplicate removal
            a = randint(15, 100)
            b = randint(15, 100)
            c = randint(15, 100)
            d = str(a) + '-' + str(b) + '-' + str(c) + '='
            if d in blend or a<b+c:
                pass
            else:
                blend.append(d)
                break

    return blend


def plus():
    p = []
    for x in range(40):
        while True:
            num1 = str(randint(62,200))
            num2 = str(randint(62,200))
            a = num1+'+'+num2+'='
            if a in p:
                pass
            else:
                p.append(a)
                break
    return p

def subtract():
    s = []
    for x in range(40):
        while True:
            a = randint(64,210)
            b = randint(64,210)
            if a>b:
                c = str(a)+'-'+str(b)+'='
                if c in s:
                    pass
                else:
                    s.append(c)
                    break
    return s

def printFile(filename):
    import tempfile
    import win32api
    import win32print
    open(filename, "r")
    win32api.ShellExecute(
        0,
        "print",
        filename,
        '/d:"%s"' % win32print.GetDefaultPrinter(),
        ".",
        0
    )

def work(fileName,subject,heading='今日数学题'):
    cols = 2
    all_rows = 40
    try:
        d = Document(fileName)
    except:
        d = Document()
    d.add_heading(heading,0)
    section = d.sections[0]
    section.left_margin = Cm(2)
    table = d.add_table(rows=1,cols=cols)
    out = subject
    for x in range(len(out)):
        out[x] = '('+str(x+1)+')  '+out[x]
    text = table.rows[0].cells
    text[0].text = out[0]
    text[1].text = out[1]
    for x in range(cols, all_rows, cols):
        row_cells = table.add_row().cells
        row_cells[0].text = out[x]
        row_cells[1].text = out[x+1]

    #make the word biger
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(14)

    d.save(fileName)

try:
    os.remove('work.docx')
    os.remove('answer.docx')
except:
    pass
for x in range(7):
    subject = blend()
    answer(days=x + 1, subject=subject)
    work('work.docx',subject=subject,heading='第'+an2cn(x+1)+'天')

printFile('work.docx')







